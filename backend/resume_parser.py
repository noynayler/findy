"""Resume text extraction from PDF and Word (.docx; optional .doc via pandoc)."""

import io
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple

import pdfplumber


_pytesseract = None
_convert_from_bytes = None


def _get_ocr():
    global _pytesseract, _convert_from_bytes
    if _pytesseract is None:
        try:
            import pytesseract
            from pdf2image import convert_from_bytes
            _pytesseract = pytesseract
            _convert_from_bytes = convert_from_bytes
        except ImportError:
            pass
    return _pytesseract, _convert_from_bytes


def extract_text_from_pdf(pdf_bytes: bytes, enable_ocr: bool = False, tesseract_cmd: Optional[str] = None) -> Tuple[str, bool]:
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            text_parts = [p.extract_text() for p in pdf.pages if p.extract_text()]
            if text_parts:
                text = re.sub(r"\s+", " ", "\n".join(text_parts).strip())
                text = re.sub(r"[^\w\s\n\-\.,;:!?()]", " ", text).strip()
                if len(text) > 100:
                    return text, False
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}")
    if enable_ocr:
        pytesseract, convert_from_bytes = _get_ocr()
        if pytesseract and convert_from_bytes:
            try:
                if tesseract_cmd:
                    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                images = convert_from_bytes(pdf_bytes)
                text = "\n".join(pytesseract.image_to_string(img, lang="eng") for img in images).strip()
                return re.sub(r"\s+", " ", text).strip(), True
            except Exception as e:
                print(f"OCR extraction failed: {e}")
    return "", False


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract plain text from a .docx file (Word 2007+)."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)
        text = "\n".join(parts)
        text = re.sub(r"\s+", " ", text).strip()
        return text
    except Exception as e:
        print(f"docx extraction failed: {e}")
        return ""


def extract_text_from_doc_legacy(doc_bytes: bytes) -> str:
    """
    Best-effort .doc (legacy Word) extraction using pandoc if installed.
    Returns empty string if pandoc is missing or conversion fails.
    """
    if not shutil.which("pandoc"):
        return ""
    path: Optional[str] = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            f.write(doc_bytes)
            path = f.name
        result = subprocess.run(
            ["pandoc", path, "-f", "doc", "-t", "plain"],
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
        )
        if result.returncode != 0 or not (result.stdout or "").strip():
            return ""
        return re.sub(r"\s+", " ", result.stdout).strip()
    except Exception as e:
        print(f"doc (pandoc) extraction failed: {e}")
        return ""
    finally:
        if path:
            try:
                os.unlink(path)
            except OSError:
                pass


def extract_text_from_resume(filename: str, file_bytes: bytes) -> Tuple[str, bool]:
    """
    Route by extension. Second return value is True if OCR was used (PDF only).
    """
    ext = Path(filename or "").suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes, enable_ocr=False)
    if ext == ".docx":
        t = extract_text_from_docx(file_bytes)
        return (t, False)
    if ext == ".doc":
        t = extract_text_from_doc_legacy(file_bytes)
        return (t, False)
    return "", False
